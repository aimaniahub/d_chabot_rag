"""Load text from DOCX files as pseudo-pages for the chunker."""
from __future__ import annotations

from pathlib import Path

from modules.pdf_loader import PageText
from utils.text_clean import clean_text


def load_docx_pages(path: str | Path) -> list[PageText]:
    """Extract paragraphs from a .docx into page-like units (grouped)."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX ingest. pip install python-docx"
        ) from exc

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(str(path))
    paras = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    # also tables
    for table in doc.tables:
        for row in table.rows:
            cells = [clean_text(c.text) for c in row.cells if clean_text(c.text)]
            if cells:
                paras.append(" | ".join(cells))

    if not paras:
        return []

    # Group ~1500 chars per "page" for metadata
    pages: list[PageText] = []
    buf: list[str] = []
    size = 0
    page_no = 1
    for para in paras:
        buf.append(para)
        size += len(para)
        if size >= 1500:
            pages.append(PageText(page=page_no, text="\n\n".join(buf)))
            page_no += 1
            buf = []
            size = 0
    if buf:
        pages.append(PageText(page=page_no, text="\n\n".join(buf)))
    return pages
